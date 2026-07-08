import streamlit as st
import os
import json
from pathlib import Path
from datetime import datetime
import PyPDF2
import docx
import requests
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
from anthropic import Anthropic
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Streamlit page configuration
st.set_page_config(
    page_title="RedChip Press Release Agent",
    page_icon="📰",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #c8102e 0%, #a00a25 100%);
        color: white;
        padding: 2rem;
        border-radius: 8px;
        margin-bottom: 2rem;
        text-align: center;
    }
    .section-header {
        color: #c8102e;
        border-bottom: 3px solid #c8102e;
        padding-bottom: 0.5rem;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# Helper Functions
# ─────────────────────────────────────────────────────────────────────────────

def read_pdf(file_path, label, char_limit=4000):
    """Read and extract text from PDF file."""
    try:
        reader = PyPDF2.PdfReader(file_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return f"[{label}]\n{text[:char_limit]}"
    except Exception as e:
        st.warning(f"Could not read PDF: {e}")
        return f"[{label}] Could not be read."

def read_docx(file_path, label, char_limit=4000):
    """Read and extract text from DOCX file."""
    try:
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return f"[{label}]\n{text[:char_limit]}"
    except Exception as e:
        st.warning(f"Could not read DOCX: {e}")
        return f"[{label}] Could not be read."

def read_text_file(file_path, label, char_limit=4000):
    """Read and extract text from TXT or MD file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return f"[{label}]\n{text[:char_limit]}"
    except Exception as e:
        st.warning(f"Could not read file: {e}")
        return f"[{label}] Could not be read."

def extract_section(text, start_label, end_label=None):
    """Extract a labeled section from Claude output."""
    try:
        start = text.index(start_label) + len(start_label)
        if end_label:
            end = text.index(end_label)
            return text[start:end].strip()
        else:
            return text[start:].strip()
    except ValueError:
        return "(Section not found)"

def read_client_files_from_folder(ticker, char_limit=4000):
    """Read client reference files from local ticker folder."""
    # Try multiple paths in order:
    # 1. Environment variable (for custom local paths)
    # 2. Repo's ./client-files folder (works on GitHub + Streamlit Cloud)
    # 3. Home ~/Downloads/client-files (local development)
    
    custom_path = os.getenv("CLIENT_FILES_PATH")
    if custom_path:
        client_files_dir = Path(custom_path)
    else:
        # Try repo folder first (works on Streamlit Cloud)
        repo_path = Path(__file__).parent / "client-files"
        if repo_path.exists():
            client_files_dir = repo_path
        else:
            # Fall back to home Downloads (local development)
            client_files_dir = Path.home() / "Downloads" / "client-files"
    
    ticker_folder = client_files_dir / ticker.upper()

    if not ticker_folder.exists():
        return ""

    parts = []

    for file in sorted(ticker_folder.iterdir()):
        if not file.is_file():
            continue

        file_ext = file.suffix.lower()

        try:
            if file_ext == ".pdf":
                reader = PyPDF2.PdfReader(str(file))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
                parts.append(f"[CLIENT FILE — {ticker} — {file.name}]\n{text[:char_limit]}")

            elif file_ext == ".docx":
                doc = docx.Document(str(file))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                parts.append(f"[CLIENT FILE — {ticker} — {file.name}]\n{text[:char_limit]}")

            elif file_ext in [".txt", ".md"]:
                text = file.read_text(errors="ignore")
                parts.append(f"[CLIENT FILE — {ticker} — {file.name}]\n{text[:char_limit]}")
        except Exception as e:
            st.warning(f"Could not read {file.name}: {e}")

    divider = "\n\n" + ("─" * 60) + "\n\n"
    return divider.join(parts) if parts else ""

def get_available_client_files(ticker):
    """List available client files for a given ticker."""
    # Try multiple paths in order:
    # 1. Environment variable (for custom local paths)
    # 2. Repo's ./client-files folder (works on GitHub + Streamlit Cloud)
    # 3. Home ~/Downloads/client-files (local development)
    
    custom_path = os.getenv("CLIENT_FILES_PATH")
    if custom_path:
        client_files_dir = Path(custom_path)
    else:
        # Try repo folder first (works on Streamlit Cloud)
        repo_path = Path(__file__).parent / "client-files"
        if repo_path.exists():
            client_files_dir = repo_path
        else:
            # Fall back to home Downloads (local development)
            client_files_dir = Path.home() / "Downloads" / "client-files"
    
    ticker_folder = client_files_dir / ticker.upper()

    if not ticker_folder.exists():
        return []

    return [
        f.name for f in sorted(ticker_folder.iterdir())
        if f.is_file() and f.suffix.lower() in [".pdf", ".docx", ".txt", ".md"]
    ]

def fetch_rss_feed(ticker):
    """Fetch RSS feed for a given ticker. Returns list of dicts or empty list on failure."""
    import urllib.request
    import ssl
    from html import unescape
    import re
    
    url = f"https://redchip.com/rss/company/{ticker.lower()}"
    content = None
    
    # Disable SSL verification if needed (for corporate proxies)
    ssl_context = ssl.create_default_context()
    ssl_context.check_hostname = False
    ssl_context.verify_mode = ssl.CERT_NONE
    
    # Use only urllib - skip requests to avoid 415 error
    try:
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        )
        with urllib.request.urlopen(req, timeout=30, context=ssl_context) as response:
            content = response.read()
            st.write(f"🔍 Fetched {len(content)} bytes, content-type: {response.headers.get('content-type')}")
    except Exception as e:
        st.write(f"🔍 Network error: {str(e)[:80]}")
        raise Exception(f"Could not fetch RSS: {str(e)[:80]}")

    if not content:
        raise Exception("No content downloaded")

    # Decode
    try:
        text = content.decode('utf-8')
    except:
        text = content.decode('utf-8', errors='ignore')
    
    # Check what we actually got
    st.write(f"🔍 {len(text)} chars, starts with: {repr(text[:80])}")
    
    if '<?xml' not in text and '<rss' not in text:
        st.write(f"🔍 WARNING: Response doesn't look like RSS or XML!")
        st.write(f"🔍 Full first 300 chars: {repr(text[:300])}")
        raise Exception("Response is not XML/RSS")
    
    releases = []
    
    # Strategy 1: XML parsing
    try:
        root = ET.fromstring(text.encode('utf-8'))
        channel = root.find("channel")
        items = channel.findall("item") if channel is not None else root.findall("item")
        
        st.write(f"🔍 XML parsing found {len(items)} items")
        
        if items:
            for item in items:
                title = item.findtext("title", "").strip()
                link = item.findtext("link", "").strip()
                pub_date = item.findtext("pubDate", "").strip()

                if title:
                    title = unescape(title)

                if title and link:
                    releases.append({
                        "title": title,
                        "url": link,
                        "pub_date": pub_date,
                        "ticker": ticker
                    })
            
            if releases:
                st.success(f"✅ Successfully extracted {len(releases)} releases")
                return releases
    except Exception as e:
        st.write(f"🔍 XML parsing failed: {type(e).__name__}: {str(e)[:60]}")
    
    # Strategy 2: Regex as fallback
    try:
        item_pattern = r'<item>([\s\S]*?)</item>'
        items_match = re.findall(item_pattern, text)
        
        st.write(f"🔍 Regex found {len(items_match)} items")
        
        if items_match:
            for item_text in items_match:
                title = None
                tm = re.search(r'<title>\s*<!\[CDATA\[(.*?)\]\]>\s*</title>', item_text, re.DOTALL)
                if tm:
                    title = tm.group(1).strip()
                else:
                    tm = re.search(r'<title>(.*?)</title>', item_text, re.DOTALL)
                    if tm:
                        title = tm.group(1).strip()
                
                lm = re.search(r'<link>(.*?)</link>', item_text, re.DOTALL)
                link = lm.group(1).strip() if lm else None
                
                pm = re.search(r'<pubDate>(.*?)</pubDate>', item_text, re.DOTALL)
                pub_date = pm.group(1).strip() if pm else ""
                
                if title and link:
                    title = unescape(title)
                    releases.append({
                        "title": title,
                        "url": link,
                        "pub_date": pub_date,
                        "ticker": ticker
                    })
            
            if releases:
                st.success(f"✅ Successfully extracted {len(releases)} releases (regex)")
                return releases
    except Exception as e:
        st.write(f"🔍 Regex parsing failed: {str(e)[:80]}")
    
    raise Exception("Could not parse RSS with XML or regex")

def scrape_press_release(url):
    """Scrape full press release content from URL."""
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"
    }
    try:
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
    except Exception as e:
        st.error(f"❌ Could not load press release: {e}")
        return "", "", ""

    soup = BeautifulSoup(response.text, "html.parser")

    headline = ""
    for tag in ["h1", "h2"]:
        el = soup.find(tag)
        if el:
            headline = el.get_text(strip=True)
            break

    paragraphs = []
    for p in soup.find_all("p"):
        text = p.get_text(strip=True)
        if len(text) > 60:
            paragraphs.append(text)

    full_body = "\n\n".join(paragraphs)
    words = full_body.split()
    first_500 = " ".join(words[:500])

    return headline, first_500, full_body

def generate_social_content(headline, body, url, ticker, master_examples, client_reference):
    """Generate social content using Claude."""
    api_client = Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

    prompt = f"""You are a financial communications specialist at RedChip, an investor relations and financial media firm that represents publicly traded small-cap companies.

You are drafting social media content for a new press release from one of RedChip's clients (ticker: {ticker}).

CRITICAL: Before drafting any content, study the MASTER EXAMPLES section below and understand how information flows from the original press release into each social format. Your task is not to rewrite press releases—it's to extract what matters most to investors and translate it into clear, grounded, analytical posts.

{"=" * 70}
MASTER EXAMPLES (Study the before-and-after transformation)
{"=" * 70}
{master_examples if master_examples else "No master examples provided."}

{"=" * 70}
CLIENT REFERENCE FILES FOR {ticker}
{"=" * 70}
{client_reference if client_reference else "No client reference files provided."}

{"=" * 70}
PRESS RELEASE TO CONVERT
{"=" * 70}
Headline: {headline}
Link: {url}

{body[:3000]}

{"=" * 70}
TONE OF VOICE & CONTENT STRATEGY
{"=" * 70}

Your audience: Small-cap investors, analysts, fund managers, retail and institutional investors. They're financially literate, fast-moving, and can spot hype immediately.

Your job: Take what the company announced, discern what matters most to investors, and explain why it matters, what to watch next, and what the signal is beneath the headline.

KEY GUIDING PRINCIPLES (Follow these in order):

1. LEAD WITH TICKER AND THE NEWS
   - Every post opens with Company Name $TICKER: and a specific, high-impact headline
   - Investors must know immediately who this is and what happened
   - No generic announcement language

2. RANK THE NEWS, THEN LEAD WITH WHAT RANKS HIGHEST
   - Prioritize in this order: quantified financial impact (revenue, margin, deal size) > 
     material catalysts (partnerships, approvals, contracts) > substantiated superlatives 
     (first, largest, only) > strategic commentary > company boilerplate
   - The highest-ranking item anchors the headline
   - Boilerplate never appears in social posts

3. SAY IT ONCE
   - Every fact appears in exactly one place—either headline or bullets, never both (even reworded)
   - If a bullet just restates the headline stat, cut it and replace with unused information, or drop it
   - No redundancy

4. TRANSLATE, DON'T TRANSCRIBE
   - Don't just restate the press release
   - Restate the news in plain language, then explain what it means for the company's trajectory
   - Address: revenue potential, market position, regulatory pathway, execution progress

5. MAKE THE NUMBERS DO THE WORK
   - Specific figures build credibility: $2.8B project value, 20,736 GPUs, 80 MW capacity
   - Concrete numbers anchor the story and give investors something to evaluate
   - Never use vague intensifiers like "significant," "substantial," "major"

6. STRUCTURE FOR SCANNERS
   - Investor audiences skim
   - Use "Why Investors Should Be Watching" or "Key Investor Takeaways" with bullet points
   - Surface most important details fast
   - Each bullet should be distinct (no overlap with headline)

7. CONFIDENT, NOT PROMOTIONAL
   - No hype. No exclamation points to manufacture excitement
   - No vague language: "exciting opportunity," "game-changing," "revolutionary," "innovative"
   - News speaks for itself. Your job is to frame it clearly
   - Tone: analytical, informed, matter-of-fact

8. MATCH THE PLATFORM
   - X (Twitter): Punchy analytical take. Tight paragraphs. Clear so-what. Specific bullets.
   - LinkedIn/Instagram/Facebook: Fuller breakdown. Emojis as section markers (not decoration). 
     Bullet points. Always include link to full release.

{"=" * 70}
OUTPUT FORMAT
{"=" * 70}

Produce exactly eight outputs below. Use ONLY these headers (nothing else between them):

GRAPHIC TITLE 1:
[A sharp, specific headline for a social media graphic. Lead with company name and what happened. 
Surface the most compelling number or milestone. No period. One sentence.]

GRAPHIC SUBTITLE 1:
[One sentence expanding on the headline—adds context, a key figure, or the investor angle. 
No period.]

GRAPHIC TITLE 2:
[A different angle on the same news—reframe around a different number, milestone, or implication. 
One sentence, no period.]

GRAPHIC SUBTITLE 2:
[One sentence expanding on Title 2. No period.]

X POST 1:
[Long-form analytical post for Twitter/X.
- Open with: Company Name $TICKER: [Compelling Headline]
- Tight analytical paragraphs (lead with news, then the so-what)
- Surface most important number in first or second paragraph
- Close with forward-looking investor takeaway
- Include "Why Investors Should Be Watching:" with bullet points for complex announcements
- No emojis. No exclamation points
- Match the length and depth of X examples in the reference guide]

X POST 2:
[Long-form analytical post from a different angle—same format as Post 1, but different lead, 
different framing, different takeaway]

LINKEDIN/INSTAGRAM/FACEBOOK POST 1:
[LinkedIn/Facebook/Instagram post.
- Open with: relevant emoji + Company Name (Exchange: TICKER): [Compelling Headline]
- Full breakdown: what happened, why it matters, what investors should watch
- Use emojis as section markers (📊 for metrics, 💡 for takeaway, 🔍 for analysis, etc.)
- Include "Why investors should be paying attention:" or "Key Investor Takeaways:" with bullet points
- Surface key figures and milestones in bullets
- Close with: Read the full PR: {url}
- End with 5-10 relevant hashtags (#TICKER, #Sector, #Investing, etc.)
- No period on section headers
- Length: 200-400 words]

LINKEDIN/INSTAGRAM/FACEBOOK POST 2:
[Same format as Post 1—different angle, different lead emoji, different hook, different bullet 
point emphasis]

IMPORTANT CHECKLIST BEFORE SUBMITTING:
✓ Ticker symbol in first line of each post
✓ Lead fact is highest-ranking (financial > catalyst > superlative > commentary)
✓ No boilerplate ("pleased to announce," "excited," etc.)
✓ Every number is concrete (no "significant," "substantial," "leading")
✓ No exclamation points
✓ No vague language ("game-changing," "innovative," "exciting opportunity")
✓ Each fact appears only once (headline or bullets, never both)
✓ Platform formatting followed exactly
✓ Clear "so-what" (what does this mean for the company's trajectory?)
✓ Link to full release included
"""

    message = api_client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text

def display_results(result):
    """Display formatted results."""
    raw = result["social_content"]
    sections = {
        "gt1": extract_section(raw, "GRAPHIC TITLE 1:", "GRAPHIC SUBTITLE 1:"),
        "gs1": extract_section(raw, "GRAPHIC SUBTITLE 1:", "GRAPHIC TITLE 2:"),
        "gt2": extract_section(raw, "GRAPHIC TITLE 2:", "GRAPHIC SUBTITLE 2:"),
        "gs2": extract_section(raw, "GRAPHIC SUBTITLE 2:", "X POST 1:"),
        "x1": extract_section(raw, "X POST 1:", "X POST 2:"),
        "x2": extract_section(raw, "X POST 2:", "LINKEDIN/INSTAGRAM/FACEBOOK POST 1:"),
        "li1": extract_section(raw, "LINKEDIN/INSTAGRAM/FACEBOOK POST 1:", "LINKEDIN/INSTAGRAM/FACEBOOK POST 2:"),
        "li2": extract_section(raw, "LINKEDIN/INSTAGRAM/FACEBOOK POST 2:", None),
    }

    st.markdown(f"### 📰 {result['headline']}")
    st.markdown(f"**Ticker:** `{result['ticker']}` | **Link:** [{result['url'][:50]}...]({result['url']})")

    with st.expander("📋 Press Release Summary (First 500 words)", expanded=False):
        st.write(result['summary'])

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### 🎨 Graphic 1")
        st.markdown(f"**{sections['gt1']}**")
        st.info(sections['gs1'])

    with col2:
        st.markdown("#### 🎨 Graphic 2")
        st.markdown(f"**{sections['gt2']}**")
        st.info(sections['gs2'])

    st.markdown("#### 𝕏 Twitter / X Posts")

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Post 1**")
        with st.container(border=True):
            st.text(sections['x1'])

    with col2:
        st.markdown("**Post 2**")
        with st.container(border=True):
            st.text(sections['x2'])

    st.markdown("#### 💼 LinkedIn / Instagram / Facebook")

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Post 1**")
        with st.container(border=True):
            st.text(sections['li1'])

    with col2:
        st.markdown("**Post 2**")
        with st.container(border=True):
            st.text(sections['li2'])


# ─────────────────────────────────────────────────────────────────────────────
# Main App
# ─────────────────────────────────────────────────────────────────────────────

st.markdown("""
<div class="main-header">
  <h1>📰 RedChip Press Release Agent</h1>
  <p>Generate professional social content from press releases</p>
</div>
""", unsafe_allow_html=True)

# Sidebar configuration
with st.sidebar:
    st.header("⚙️ Configuration")

    api_key = st.text_input(
        "Anthropic API Key",
        value=os.getenv("ANTHROPIC_API_KEY", ""),
        type="password",
        help="Get your key from console.anthropic.com"
    )

    if api_key:
        os.environ["ANTHROPIC_API_KEY"] = api_key

    st.divider()

    st.markdown("**💡 How It Works:**")
    st.markdown("""
1. Enter ticker symbol
2. Add reference docs to `client-files/{TICKER}/`
3. Fetch a press release from RSS
4. Get 8 formatted social posts
    """)

# Pre-load master guide once (used both for display and generation)
master_guide_path = Path(__file__).parent / "master_guide.md"
if master_guide_path.exists():
    master_guide_text = master_guide_path.read_text()
else:
    master_guide_text = ""

# Main content
tab1, tab2 = st.tabs(["📝 Generate Content", "📚 Reference Guide"])

with tab1:
    col1, col2 = st.columns([2, 1])

    with col1:
        st.markdown("### Step 1: Ticker Information")
        ticker = st.text_input(
            "Company Ticker Symbol",
            placeholder="e.g., GLYFF, CYBN",
            help="The stock ticker for the company"
        ).upper()

    with col2:
        st.markdown("### ")
        if ticker:
            st.success(f"Ticker: `{ticker}`")

    st.divider()

    st.markdown("### Step 2: Client Reference Documents")

    if ticker:
        client_files = get_available_client_files(ticker)
        if client_files:
            st.success(f"✅ Found {len(client_files)} reference file(s) for {ticker}")
            with st.expander("📁 View loaded files"):
                for f in client_files:
                    st.markdown(f"- `{f}`")
        else:
            st.info(
                f"ℹ️ No reference files found for **{ticker}**. "
                f"Reference files are pulled from the app repository."
            )
    else:
        st.info("👆 Enter a ticker in Step 1 to check for reference documents")

    st.divider()
    st.markdown("#### 📚 Master Social Content Reference Guide")
    st.caption("Automatically loaded from master_guide.md — RedChip professional tone and examples")

    if master_guide_text:
        st.success(f"✅ Loaded {len(master_guide_text)} characters of professional guidance")
        with st.expander("📖 View Master Guide Reference"):
            st.markdown(master_guide_text[:1000] + "...\n\n*(Full guide loaded automatically)*")
    else:
        st.warning("⚠️ master_guide.md not found — using default tone")

    st.divider()

    st.markdown("### Step 3: Press Release — Fetch from RSS or Paste Manually")

    if not ticker:
        st.info("👆 Enter a ticker symbol above to fetch press releases")
    else:
        # RSS Feed Option
        st.markdown("#### 📡 Option A: Fetch from RSS Feed")
        if st.button("📡 Fetch Latest Press Releases", use_container_width=True):
            st.write(f"🔍 **Debug:** Starting fetch for {ticker}")
            with st.spinner(f"📡 Fetching RSS feed for {ticker}..."):
                try:
                    fetched = fetch_rss_feed(ticker)
                    st.write(f"🔍 **Debug:** Got {len(fetched)} items back")
                except Exception as debug_e:
                    st.write(f"🔍 **Debug Error in fetch:** {type(debug_e).__name__}: {str(debug_e)}")
                    fetched = []
            
            if fetched:
                st.session_state.releases = fetched
                st.session_state.releases_ticker = ticker
                st.session_state.expanded_release = 0
                st.write(f"🔍 **Debug:** Saved {len(fetched)} releases to session")
            else:
                st.session_state.releases = []
                st.write(f"🔍 **Debug:** Fetch returned empty list")

        # Always render releases from session state so buttons survive reruns
        releases = st.session_state.get("releases", [])
        if releases and st.session_state.get("releases_ticker") == ticker:
            st.success(f"✅ Found {len(releases)} press release(s)")
            st.divider()

            for idx, release in enumerate(releases):
                is_expanded = st.session_state.get("expanded_release", 0) == idx
                title_display = release['title'][:70] + ("..." if len(release['title']) > 70 else "")
                pub_date_display = release['pub_date'][:10] if release['pub_date'] else ""

                with st.expander(f"📰 {title_display} ({pub_date_display})", expanded=is_expanded):
                    st.markdown(f"**Title:** {release['title']}")
                    st.markdown(f"**URL:** [{release['url']}]({release['url']})")
                    st.markdown(f"**Published:** {release['pub_date']}")

                    col1, col2 = st.columns([3, 1])

                    with col1:
                        if st.button("✨ Generate Social Content", use_container_width=True, type="primary", key=f"generate_tab1_{idx}"):
                            st.session_state.expanded_release = idx
                            if not api_key:
                                st.error("❌ Anthropic API Key is required (paste in sidebar)")
                            else:
                                with st.spinner("📄 Scraping press release..."):
                                    pr_headline, pr_summary, pr_body = scrape_press_release(release['url'])

                                if not pr_headline:
                                    st.error("❌ Could not extract headline from press release")
                                else:
                                    st.success(f"✅ Scraped: {pr_headline[:80]}...")

                                    with st.spinner("🔄 Loading reference documents..."):
                                        client_reference_text = read_client_files_from_folder(ticker)

                                    with st.spinner("🤖 Generating social content with Claude..."):
                                        try:
                                            social_content = generate_social_content(
                                                pr_headline,
                                                pr_body,
                                                release['url'],
                                                ticker,
                                                master_guide_text,
                                                client_reference_text
                                            )

                                            st.session_state.last_result = {
                                                "ticker": ticker,
                                                "headline": pr_headline,
                                                "url": release['url'],
                                                "summary": pr_summary,
                                                "social_content": social_content,
                                                "source": "RSS Feed"
                                            }
                                            st.success("✅ Content generated!")

                                        except Exception as e:
                                            st.error(f"❌ Error generating content: {str(e)}")

                    with col2:
                        if st.button("🔄 Check", use_container_width=True, key=f"check_tab1_{idx}"):
                            st.session_state.expanded_release = idx
                            with st.spinner("🔄 Checking..."):
                                pr_headline, pr_summary, pr_body = scrape_press_release(release['url'])

                            if pr_headline:
                                st.success(f"✅ Valid: {pr_headline[:50]}...")
                            else:
                                st.warning("⚠️ Could not scrape this release")
        
        # Manual Input Option
        st.markdown("#### ✏️ Option B: Paste Press Release Manually")
        st.info("If RSS isn't available, paste the details below:")
        
        col1, col2 = st.columns(2)
        with col1:
            manual_title = st.text_input("Press Release Title", placeholder="e.g., Company Announces Q1 Results", key="manual_title")
        with col2:
            manual_url = st.text_input("Press Release URL", placeholder="https://example.com/press-release", key="manual_url")
        
        manual_body = st.text_area(
            "Press Release Content (first 3000 chars will be used)",
            placeholder="Paste the full press release text here...",
            height=150,
            key="manual_body"
        )
        
        if st.button("✨ Generate from Manual Input", use_container_width=True, type="primary", key="generate_manual"):
            if not ticker or not manual_title or not manual_url or not manual_body:
                st.error("❌ Please fill in all fields: Ticker, Title, URL, and Content")
            elif not api_key:
                st.error("❌ Anthropic API Key is required (paste in sidebar)")
            else:
                with st.spinner("🔄 Loading reference documents..."):
                    client_reference_text = read_client_files_from_folder(ticker)

                with st.spinner("🤖 Generating social content with Claude..."):
                    try:
                        social_content = generate_social_content(
                            manual_title,
                            manual_body,
                            manual_url,
                            ticker,
                            master_guide_text,
                            client_reference_text
                        )

                        st.session_state.last_result = {
                            "ticker": ticker,
                            "headline": manual_title,
                            "url": manual_url,
                            "summary": manual_body[:500],
                            "social_content": social_content,
                            "source": "Manual Input"
                        }
                        st.success("✅ Content generated!")

                    except Exception as e:
                        st.error(f"❌ Error generating content: {str(e)}")

    if hasattr(st.session_state, 'last_result'):
        st.divider()
        st.markdown("### 📊 Generated Content")
        display_results(st.session_state.last_result)

        st.divider()
        result_json = json.dumps(st.session_state.last_result, indent=2)
        st.download_button(
            "📥 Download as JSON",
            result_json,
            file_name=f"redchip_{st.session_state.last_result['ticker']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json"
        )


with tab2:
    st.markdown("### 📚 Setup Instructions")
    st.info("""
**Quick Start:**

1. Get your **Anthropic API Key** from [console.anthropic.com](https://console.anthropic.com)
2. Paste it in the Configuration panel on the left
3. Add client reference docs to `client-files/{TICKER}/` (factsheets, investor decks, etc.)
4. Enter a ticker and either:
   - **Fetch from RSS** — if available for your ticker
   - **Paste manually** — if RSS isn't available

The AI will produce:
- 2 Graphic titles + subtitles (for social media graphics)
- 2 X/Twitter analytical posts
- 2 LinkedIn/Instagram/Facebook posts
    """)

    st.markdown("### 📡 About RSS Feeds")
    st.warning("""
**Current Status:** RedChip's RSS feed may have limited availability or changed formats.

**If RSS isn't working:**
- ✅ Use the **Manual Input** option to paste press releases directly
- ✅ The manual input is just as powerful — you control exactly what content is analyzed
- 💡 Check RedChip's website directly for recent releases
- 💡 Copy-paste the full press release text for best results

**If RSS works:**
- The feed will auto-populate recent releases for your ticker
- Click "Generate Social Content" to process any release
    """)

    st.markdown("### 🎯 Reference Document Best Practices")
    st.markdown("""
**For Master Examples / Social Content Reference Guide (`master_guide.md`):**
- Include RedChip's tone of voice guidelines
- Add real examples of approved posts
- Show formatting and emoji usage patterns
- Clarify what to avoid (hype, vague language, etc.)

**For Client Reference Files (`client-files/{TICKER}/`):**
- Factsheets with company overview and metrics
- Investor decks with key talking points
- Previous press releases (for tone consistency)
- Industry/regulatory context
    """)

    st.markdown("### 📋 Troubleshooting")
    st.markdown("""
**Content not generating?**
- Check your API key is valid
- Ensure the press release URL is accessible (for RSS) or paste full text (for manual)
- If using RSS, try the manual input option instead

**RSS fetch failing?**
- This is expected — use the manual input option instead
- Manual input gives you more control over what content is analyzed anyway

**Quality not matching your style?**
- Add a more detailed `master_guide.md` with real approved examples
- Include more client reference documents in `client-files/{TICKER}/`
- Ensure the press release content is comprehensive

**Need to make changes?**
- Generate new content with updated documents or a different release
- Each generation is independent
    """)
